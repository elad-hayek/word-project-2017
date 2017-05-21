# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
import os
import docx2txt
import mail_sending

RED = '\033[91m'   # red
GREEN = '\033[92m'  #green
END = '\033[0m'    # reset

error_messages = {
    'save doc': (RED + 'ERROR: path does not exist, action terminated' + END),
    'add new picture': (RED + 'ERROR: path does not exist, action terminated' + END),
    'write to table': ((RED + 'ERROR: there is no such row number, action terminated' + END),
                      (RED + 'ERROR: there is no such column number, action terminated' + END),
                      (RED + 'ERROR: there is no table to edit, action terminated' + END)),
    'add new paragraph': ((RED + 'ERROR: there is no such color, action terminated' + END),
                          (RED + 'ERROR: there is not such style, action terminated' + END)),
    'import existing doc': (RED + 'ERROR: path does not exist, action terminated' + END),
    'return to last save': (RED + 'WARNING: your document was not previously saved, action terminated' + END),
    'email doc': (RED + 'ERROR: your document was not previously saved, action terminated' + END)}
confirm_messages = {
    'save doc': (GREEN + 'document saved successfully' + END),
    'add new picture': (GREEN + 'picture added successfully' + END),
    'add new table': (GREEN + 'table added successfully' + END),
    'write to table': (GREEN + 'successfully writen to the table' + END),
    'add new paragraph': (GREEN +  'paragraph added successfully' + END),
    'import existing doc': (GREEN + 'successfully imported a document' + END),
    'add new heading': (GREEN + 'heading added successfully' + END),
    'add new page break': (GREEN + 'new page break added successfully' + END),
    'return to last save': (GREEN + 'WARNING: your document was not previously saved, action terminated' + END),
    'email doc': (GREEN + 'your email has been sent' + END)}

colors = {'R': (0xff, 0x0, 0x0), 'G': (0x0, 0xff, 0x0), 'b': (0x0, 0x0, 0xff), 'B': (0x0, 0x0, 0x0)} #red, green, blue, black


TABLE_STYLE = 'Table Grid'




class Documents():
    def __init__(self):
        self.__path = ''
        self.__document = Document()
        self.__document.save('beta.docx')
        self.__table = None
        self.__name = ''


    def save_doc(self, args_list):
        path = args_list[0]
        name = args_list[1]
        if not os.path.exists(path):
            return error_messages['save doc']
        self.__path = path + '\\' + name + '.docx'
        print 'saving document'
        print path
        self.__document.save(path + '\\' + name + '.docx')
        os.remove('beta.docx')
        self.__name = name
        return confirm_messages['save doc']


    def add_new_picture(self, args_list):
        path = args_list[0]
        width = args_list[1]
        if not os.path.exists(path):
            return error_messages['add new picture']
        self.__document.add_picture(path, width=Inches(float(width)))
        return confirm_messages['add new picture']

    def add_new_table(self, args_list):
        rows = args_list[0]
        lines = args_list[1]
        table = self.__document.add_table(rows=int(rows), cols=int(lines))
        table.style = TABLE_STYLE
        self.__table = table
        return confirm_messages['add new table']



    def write_to_table(self, args_list):
        row = args_list[0]
        column = args_list[1]
        text = args_list[2]
        print row+' '+column, 'raw, column'
        if not self.__table:
            return error_messages['write to table'][2]
        if len(self.__table.rows) < int(row):
            return error_messages['write to table'][0]
        if len(self.__table.columns) < int(column):
            return error_messages['write to table'][1]
        cell = self.__table.cell(int(row), int(column))
        cell.text = text
        return confirm_messages['write to table']

    def add_new_paragraph(self, args_list):
        text = args_list[0]
        style = args_list[1]
        color = args_list[2]
        if color not in colors:
            return error_messages['add new paragraph'][0]
        size = args_list[3]
        font_name = args_list[4]
        run = self.__document.add_paragraph().add_run(text)
        font = run.font
        font.name = font_name
        font.color.rgb = RGBColor(colors[color][0], colors[color][1], colors[color][2])
        font.size = Pt(int(size))
        if style == 'B':
            font.bold = True
        elif style == 'I':
            font.italic = True
        elif style == 'N':
            pass
        else:
            return error_messages['add new paragraph'][1]
        return confirm_messages['add new paragraph']



    def import_existing_doc(self, args_list):
        path = args_list[0]
        if not os.path.exists(path):
            return error_messages['import existing doc']
        self.__document = Document(path)
        return confirm_messages['import existing doc']

    def add_new_heading(self, args_list):
        text = args_list[0]
        self.__document.add_heading(text, level=0)
        return confirm_messages['add new heading']

    def add_new_page_break(self, args_list):
        print 'adding new page break'
        self.__document.add_page_break()
        return confirm_messages['add new page break']

    def return_to_last_save(self, args_list):
        if os.path.exists('beta.docx'):
             return error_messages['return to last save']
        self.__document = Document(self.__path)
        return confirm_messages['return to last save']

    def live_docx(self, args_list):
        if os.path.exists('beta.docx'):
            self.__document.save('beta.docx')
            live_view = docx2txt.process('beta.docx')
        else:
            live_view = docx2txt.process(self.__path)
        return GREEN + 'THIS IS YOUR FILE TEXT:\n'+str(live_view)+'\n'+END


    def email_doc(self, args_list):
        if os.path.exists('beta.docx'):
            return error_messages['email doc']
        from_email = args_list[0]
        password = args_list[1]
        to_email = args_list[2]
        mail_sending.send(from_email, password, to_email, self.__path, self.__name)
        return confirm_messages['email doc']




