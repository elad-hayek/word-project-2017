# -*- coding: utf-8 -*-
"""
Description      : The class where all the actions related to the document
                    happens

Author           : Elad Hayek
FileName         : client.py
Date             : 27.5.17
Version          : 1.0
"""

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
import os
import mail_sending
from colorama import Fore
import subprocess


TABLE_STYLE = 'Table Grid'
CMD_WORD_PROCESS = 'C:\Program Files\Microsoft Office\Office14\WINWORD.EXE'
error_messages = {
    'save doc': (
        Fore.LIGHTRED_EX, 'ERROR: path does not exist, action terminated'),
    'add new picture': (
        Fore.LIGHTRED_EX, 'ERROR: path does not exist, action terminated'),
    'write to table': (
        (Fore.LIGHTRED_EX,
         'ERROR: there is no such row number, action terminated'),
        (Fore.LIGHTRED_EX,
         'ERROR: there is no such column number, action terminated'),
        (Fore.LIGHTRED_EX,
         'ERROR: there is no table to edit, action terminated')),
    'add new paragraph': (
        (Fore.LIGHTRED_EX,
         'ERROR: there is no such color, action terminated'),
        (Fore.LIGHTRED_EX,
         'ERROR: there is not such style, action terminated'),
        (Fore.LIGHTRED_EX,
         'ERROR: the font size is invalid, action terminated')),
    'import existing doc': (
        Fore.LIGHTRED_EX, 'ERROR: path does not exist, action terminated'),
    'return to last save': (
        Fore.LIGHTRED_EX,
        'WARNING: your document was not previously saved, action terminated'),
    'email doc': (
        Fore.LIGHTRED_EX,
        'ERROR: your document was not previously saved, action terminated')}

confirm_messages = {
    'save doc': (
        Fore.LIGHTGREEN_EX, 'document saved successfully'),
    'add new picture': (
        Fore.LIGHTGREEN_EX, 'picture added successfully'),
    'add new table': (
        Fore.LIGHTGREEN_EX, 'table added successfully'),
    'write to table': (
        Fore.LIGHTGREEN_EX, 'successfully writen to the table'),
    'add new paragraph': (
        Fore.LIGHTGREEN_EX,  'paragraph added successfully'),
    'import existing doc': (
        Fore.LIGHTGREEN_EX, 'successfully imported a document'),
    'add new heading': (
        Fore.LIGHTGREEN_EX, 'heading added successfully'),
    'add new page break': (
        Fore.LIGHTGREEN_EX, 'new page break added successfully'),
    'email doc': (
        Fore.LIGHTGREEN_EX, 'your email has been sent'),
    'return to last save': (
        Fore.LIGHTGREEN_EX, 'successfully returned to last save'),
    'live docx': (
        Fore.LIGHTGREEN_EX, 'successfully opend your docx file')}

colors = {'R': (0xff, 0x0, 0x0), 'G': (0x0, 0xff, 0x0), 'b': (0x0, 0x0, 0xff),
          'B': (0x0, 0x0, 0x0)}  # red, green, blue, black


class Documents():
    def __init__(self):
        """
        creates a new document
        """
        self.__path = ''
        self.__document = Document()
        self.__document.save('beta.docx')
        self.__table = None
        self.__name = ''

    def save_doc(self, args_list):
        """
        saves the document according to the path the client sent

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        path = args_list[0]
        name = args_list[1]
        if not os.path.exists(path):
            return error_messages['save doc']
        self.__path = path + '\\' + name + '.docx'
        print 'saving document'
        print path
        self.__document.save(path + '\\' + name + '.docx')
        if os.path.exists('beta.docx'):
            os.remove('beta.docx')
        self.__name = name
        return confirm_messages['save doc']

    def add_new_picture(self, args_list):
        """
        adds a new picture to the document according to the
        path the client sent

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        path = args_list[0]
        width = args_list[1]
        if not os.path.exists(path):
            return error_messages['add new picture']
        self.__document.add_picture(path, width=Inches(float(width)))
        return confirm_messages['add new picture']

    def add_new_table(self, args_list):
        """
        adds a new table to the document according to the number of rows
        and columns the client sent

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        rows = args_list[0]
        lines = args_list[1]
        table = self.__document.add_table(rows=int(rows), cols=int(lines))
        table.style = TABLE_STYLE
        self.__table = table
        return confirm_messages['add new table']

    def write_to_table(self, args_list):
        """
        writes to a specific cell in the last created table according to the
        cell index the client sent

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        row = args_list[0]
        column = args_list[1]
        text = args_list[2]
        print len(self.__table.rows), 'table rows'
        print len(self.__table.columns), 'table colums'
        if not self.__table:
            return error_messages['write to table'][2]
        if len(self.__table.rows) <= int(row):
            return error_messages['write to table'][0]
        if len(self.__table.columns) <= int(column):
            return error_messages['write to table'][1]
        cell = self.__table.cell(int(row), int(column))
        cell.text = text
        return confirm_messages['write to table']

    def add_new_paragraph(self, args_list):
        """
        adds a new paragraph to the document according to the client's request

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        text = args_list[0]
        style = args_list[1]
        color = args_list[2]
        if color not in colors:
            return error_messages['add new paragraph'][0]
        size = args_list[3]
        if not 1 <= int(size) <= 72:
            print 'wrong size'
            return error_messages['add new paragraph'][2]
        font_name = args_list[4]
        run = self.__document.add_paragraph().add_run(text)
        font = run.font
        font.name = font_name
        font.color.rgb = RGBColor(colors[color][0], colors[color][1],
                                  colors[color][2])
        font.size = Pt(int(size))
        if style == 'B':
            font.bold = True
        elif style == 'I':
            font.italic = True
        elif style == 'N':
            pass
        else:
            return error_messages['add new paragraph'][1]
        return confirm_messages['add new paragraph']

    def import_existing_doc(self, args_list):
        """
        importing a document according to the path the client sent

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        path = args_list[0]
        if not os.path.exists(path):
            return error_messages['import existing doc']
        self.__document = Document(path)
        return confirm_messages['import existing doc']

    def add_new_heading(self, args_list):
        """
        adds a new heading to the document according to the client's request

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        text = args_list[0]
        self.__document.add_heading(text, level=0)
        return confirm_messages['add new heading']

    def add_new_page_break(self, args_list):
        """
        adds a new page break to the document

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        print 'adding new page break'
        self.__document.add_page_break()
        return confirm_messages['add new page break']

    def return_to_last_save(self, args_list):
        """
        returns the document to its last save

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        if os.path.exists('beta.docx'):
            return error_messages['return to last save']
        self.__document = Document(self.__path)
        return confirm_messages['return to last save']

    def live_docx(self, args_list):
        """
        saves and opens the document

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        if os.path.exists('beta.docx'):
            self.__document.save('beta.docx')
            cmd_file_path = 'beta.docx'
            cmd = [CMD_WORD_PROCESS, cmd_file_path]
            process = subprocess.Popen(cmd, stdout=subprocess.PIPE)
            process.wait()
        else:
            self.__document.save(self.__path)
            cmd = [CMD_WORD_PROCESS, self.__path]
            process = subprocess.Popen(cmd, stdout=subprocess.PIPE)
            process.wait()
        return confirm_messages['live docx']

    def email_doc(self, args_list):
        """
        emails the document to a gmail account

        :arg args_list = the arguments the client sent
        :type args_list = list
        """
        if os.path.exists('beta.docx'):
            return error_messages['email doc']
        to_email = args_list[0]
        mail_sending.send(to_email, self.__path, self.__name)
        return confirm_messages['email doc']
