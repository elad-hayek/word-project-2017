# -*- coding: utf-8 -*-
from docx import Document
from docx import table

TABLE_STYLE = 'Table Grid'
class Documents():
    def __init__(self, path=''):
        self.__path = path
        self.__document = Document()


    def save_doc(self, args_list):
        path = args_list[0]
        print 'saving document'
        print path
        self.__document.save(path)

    def add_new_picture(self, args_list):
        width = args_list[0]
        pass

    def add_new_table(self, args_list):
        rows = args_list[0]
        lines = args_list[1]
        table = self.__document.add_table(rows=int(rows), cols=int(lines))
        table.style = TABLE_STYLE


    def write_to_table(self, args_list):
        row = args_list[0]
        column = args_list[1]
        text = args_list[2]
        style = args_list[3]
        color = args_list[4]
        pass

    def add_new_paragraph(self, args_list):
        text = args_list[0]
        style = args_list[1]
        color = args_list[2]
        pass

    def add_new_word(self, args_list):
        text = args_list[0]
        style = args_list[1]
        color = args_list[2]
        return self

    def import_existing_doc(self, args_list):
        path = args_list[0]
        self.__document = Document(path)

    def add_new_heading(self, args_list):
        text = args_list[0]
        style = args_list[1]
        color = args_list[2]

    def add_new_page_break(self, args_list):
        print 'adding new page break'
        self.__document.add_page_break()

    def email_doc(self, args_list):
        from_email = args_list[0]
        to_email = args_list[1]
        pass



