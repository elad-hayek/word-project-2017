# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

TABLE_STYLE = 'Table Grid'

colors = {'red': (0xff, 0x0, 0x0), 'green': (0x0, 0xff, 0x0), 'blue': (0x0, 0x0, 0xff), 'black': (0x0, 0x0, 0x0)}
class Documents():
    def __init__(self, path=''):
        self.__path = path
        self.__document = Document()
        self.__table = []


    def save_doc(self, args_list):
        path = args_list[0]
        print 'saving document'
        print path
        self.__document.save(path)


    def add_new_picture(self, args_list):
        path = args_list[0]
        width = args_list[1]
        self.__document.add_picture(path, width=Inches(float(width)))

    def add_new_table(self, args_list):
        rows = args_list[0]
        lines = args_list[1]
        table = self.__document.add_table(rows=int(rows), cols=int(lines))
        table.style = TABLE_STYLE
        self.__table.append(table)



    def write_to_table(self, args_list):
        what_table = args_list[0]
        if what_table > len(self.__table):
            return False, 'table not exsist'
        row = args_list[1]
        column = args_list[2]
        text = args_list[3]
        cell = self.__table[int(what_table)].cell(int(row), int(column))
        cell.text = text
        return True, 'you have successfully'


    def add_new_paragraph(self, args_list):
        text = args_list[0]
        style = args_list[1]
        color = args_list[2]
        size = args_list[3]
        font_name = args_list[4]
        run = self.__document.add_paragraph().add_run(text)
        font = run.font
        font.name = font_name
        font.color.rgb = RGBColor(colors[color][0], colors[color][1], colors[color][2])
        font.size = Pt(int(size))
        if style == 'bold':
            font.bold = True
        elif style == 'italic':
            font.italic = True



    def import_existing_doc(self, args_list):
        path = args_list[0]
        self.__document = Document(path)

    def add_new_heading(self, args_list):
        text = args_list[0]
        style = args_list[1]
        color = args_list[2]

    def add_new_page_break(self, args_list):
        print 'adding new page break'
        self.__document.add_page_break()

    def email_doc(self, args_list):
        from_email = args_list[0]
        to_email = args_list[1]
        pass



